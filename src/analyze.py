"""Core contract analysis module using Claude API."""

import os
import json
import base64
import re
from pathlib import Path
from typing import Optional, Literal
from datetime import datetime
import anthropic

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load prompts and knowledge
PROMPTS_DIR = Path(__file__).parent / "prompts"
KNOWLEDGE_DIR = Path(__file__).parent / "knowledge"
JURISDICTIONS_DIR = KNOWLEDGE_DIR / "jurisdictions"


def load_prompt(name: str) -> str:
    """Load a prompt template from the prompts directory."""
    prompt_path = PROMPTS_DIR / f"{name}.md"
    if prompt_path.exists():
        return prompt_path.read_text(encoding="utf-8")
    raise FileNotFoundError(f"Prompt template not found: {name}")


def load_risk_patterns() -> dict:
    """Load risk patterns from knowledge base."""
    patterns_path = KNOWLEDGE_DIR / "risk_patterns.json"
    if patterns_path.exists():
        return json.loads(patterns_path.read_text(encoding="utf-8"))
    return {"patterns": []}


def load_jurisdiction_knowledge(jurisdiction: str, contract_type: str = "employment") -> dict:
    """
    Load jurisdiction-specific legal knowledge.
    
    Args:
        jurisdiction: Country code (e.g., "us", "china", "eu")
        contract_type: Type of contract (e.g., "employment", "nda")
        
    Returns:
        dict with jurisdiction-specific knowledge
    """
    knowledge_path = JURISDICTIONS_DIR / jurisdiction / f"{contract_type}.json"
    if knowledge_path.exists():
        return json.loads(knowledge_path.read_text(encoding="utf-8"))
    return {}


def get_available_jurisdictions() -> list:
    """Get list of available jurisdiction knowledge bases."""
    if not JURISDICTIONS_DIR.exists():
        return []
    return [d.name for d in JURISDICTIONS_DIR.iterdir() if d.is_dir()]


def build_jurisdiction_context(jurisdiction: str, contract_type: str = "employment") -> str:
    """
    Build a context string from jurisdiction knowledge for prompt injection.
    
    Args:
        jurisdiction: Country code
        contract_type: Type of contract
        
    Returns:
        Formatted string with relevant legal knowledge
    """
    knowledge = load_jurisdiction_knowledge(jurisdiction, contract_type)
    if not knowledge:
        return ""
    
    context_parts = []
    
    # Add primary laws
    if "jurisdiction" in knowledge and "primary_laws" in knowledge["jurisdiction"]:
        laws = knowledge["jurisdiction"]["primary_laws"]
        context_parts.append("## Applicable Laws")
        for law in laws:
            context_parts.append(f"- **{law['name']}** ({law['citation']}): {law['description']}")
    
    # Add jurisdiction-specific risk patterns
    if "risk_patterns" in knowledge:
        context_parts.append("\n## Jurisdiction-Specific Risk Patterns")
        for pattern in knowledge["risk_patterns"]:
            context_parts.append(f"- **{pattern['name']}** [{pattern['severity'].upper()}]: {pattern['description']}")
            if "recommendation" in pattern:
                context_parts.append(f"  - Recommendation: {pattern['recommendation']}")
    
    # Add compliance checklist
    if "compliance_checklist" in knowledge:
        checklist = knowledge["compliance_checklist"]
        if "avoid_these_clauses" in checklist:
            context_parts.append("\n## Red Flags for This Jurisdiction")
            for item in checklist["avoid_these_clauses"]:
                context_parts.append(f"- ⚠️ {item}")
    
    return "\n".join(context_parts)


def encode_pdf(pdf_path: str) -> str:
    """Encode PDF file to base64."""
    with open(pdf_path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode("utf-8")


def analyze_contract(
    pdf_path: str,
    api_key: Optional[str] = None,
    model: str = "claude-sonnet-4-20250514",
    jurisdiction: Optional[str] = None,
    contract_type: str = "employment"
) -> dict:
    """
    Analyze a contract PDF and return structured analysis.
    
    Args:
        pdf_path: Path to the PDF file
        api_key: Claude API key (defaults to ANTHROPIC_API_KEY env var)
        model: Claude model to use
        jurisdiction: Optional jurisdiction code (e.g., "us", "china", "eu")
                     If not provided, will attempt to auto-detect
        contract_type: Type of contract for jurisdiction-specific knowledge
        
    Returns:
        dict with analysis results
    """
    # Initialize client
    client = anthropic.Anthropic(api_key=api_key or os.getenv("ANTHROPIC_API_KEY"))
    
    # Load prompt and knowledge
    analysis_prompt = load_prompt("risk_analysis")
    risk_patterns = load_risk_patterns()
    
    # Load jurisdiction-specific knowledge if available
    jurisdiction_context = ""
    if jurisdiction:
        jurisdiction_context = build_jurisdiction_context(jurisdiction, contract_type)
    
    # Encode PDF
    pdf_data = encode_pdf(pdf_path)
    
    # Build the full prompt with risk patterns and jurisdiction knowledge
    jurisdiction_section = ""
    if jurisdiction_context:
        jurisdiction_section = f"""
## Jurisdiction-Specific Legal Context ({jurisdiction.upper()})

{jurisdiction_context}

When analyzing this contract, pay special attention to the jurisdiction-specific requirements and risk patterns above.
"""
    
    full_prompt = f"""{analysis_prompt}

## General Risk Patterns Reference

{json.dumps(risk_patterns, ensure_ascii=False, indent=2)}
{jurisdiction_section}
---

Please analyze the attached contract and provide a structured analysis following the format above.
If jurisdiction was not specified, first identify the likely jurisdiction from the contract content (governing law clause, party locations, legal references) and apply relevant local laws.
"""
    
    # Call Claude API with PDF
    response = client.messages.create(
        model=model,
        max_tokens=8192,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data
                        }
                    },
                    {
                        "type": "text",
                        "text": full_prompt
                    }
                ]
            }
        ]
    )
    
    # Extract the response text
    analysis_text = response.content[0].text
    
    return {
        "status": "success",
        "analysis": analysis_text,
        "model": model,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens
        }
    }


def check_completeness(
    pdf_path: str,
    api_key: Optional[str] = None,
    model: str = "claude-sonnet-4-20250514"
) -> dict:
    """
    Check contract completeness (signatures, stamps, dates).
    
    Args:
        pdf_path: Path to the PDF file
        api_key: Claude API key
        model: Claude model to use
        
    Returns:
        dict with completeness check results
    """
    client = anthropic.Anthropic(api_key=api_key or os.getenv("ANTHROPIC_API_KEY"))
    
    pdf_data = encode_pdf(pdf_path)
    
    prompt = """Check the completeness of this contract, focusing on the following aspects:

1. **Signature Check**
   - Is Party A's signature present?
   - Is Party B's signature present?
   - Are signatures in the correct location?

2. **Stamp/Seal Check** (for contracts requiring official seals)
   - Is Party A's stamp/seal present?
   - Is Party B's stamp/seal present?
   - Are stamps clear and legible?

3. **Date Check**
   - Is the signing date filled in?
   - Is the effective date specified?
   - Is the date format correct?

4. **Other Elements**
   - Is the contract number present?
   - Are page numbers complete?
   - Are there binding seals (for multi-page documents)?

Output the results in JSON format:
```json
{
  "signatures": {
    "party_a": {"present": bool, "location": "description"},
    "party_b": {"present": bool, "location": "description"}
  },
  "stamps": {
    "party_a": {"present": bool, "clear": bool},
    "party_b": {"present": bool, "clear": bool}
  },
  "dates": {
    "signing_date": {"present": bool, "value": "date or null"},
    "effective_date": {"present": bool, "value": "date or null"}
  },
  "other": {
    "contract_number": {"present": bool, "value": "number or null"},
    "page_numbers": {"present": bool, "complete": bool}
  },
  "overall_completeness": "complete/partial/incomplete",
  "issues": ["issue1", "issue2"]
}
```

IMPORTANT: Respond in the same language as the contract document.
"""
    
    response = client.messages.create(
        model=model,
        max_tokens=2048,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }
        ]
    )
    
    return {
        "status": "success",
        "result": response.content[0].text,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens
        }
    }


def compare_contracts(
    pdf_path_a: str,
    pdf_path_b: str,
    api_key: Optional[str] = None,
    model: str = "claude-sonnet-4-20250514"
) -> dict:
    """
    Compare two contracts and highlight differences.
    
    Args:
        pdf_path_a: Path to the first PDF
        pdf_path_b: Path to the second PDF
        api_key: Claude API key
        model: Claude model to use
        
    Returns:
        dict with comparison results
    """
    client = anthropic.Anthropic(api_key=api_key or os.getenv("ANTHROPIC_API_KEY"))
    
    pdf_data_a = encode_pdf(pdf_path_a)
    pdf_data_b = encode_pdf(pdf_path_b)
    
    prompt = """Compare these two contracts and analyze their differences:

## Comparison Dimensions

1. **Basic Information Differences**
   - Contract type
   - Parties involved
   - Term/Duration
   - Value/Amount

2. **Clause Differences**
   - Added clauses
   - Removed clauses
   - Modified clauses (highlight specific changes)

3. **Risk Changes**
   - Clauses with increased risk
   - Clauses with reduced risk

4. **Recommendations**
   - Should the changes be accepted?
   - Issues to be aware of

Output a structured Markdown comparison report.
Label the first document as [Contract A] and the second as [Contract B].

IMPORTANT: Respond in the same language as the contract documents.
"""
    
    response = client.messages.create(
        model=model,
        max_tokens=8192,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data_a
                        }
                    },
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data_b
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }
        ]
    )
    
    return {
        "status": "success",
        "comparison": response.content[0].text,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens
        }
    }


def extract_key_terms(
    pdf_path: str,
    api_key: Optional[str] = None,
    model: str = "claude-sonnet-4-20250514"
) -> dict:
    """
    Extract key terms from a contract in structured JSON format.
    
    Args:
        pdf_path: Path to the PDF file
        api_key: Claude API key
        model: Claude model to use
        
    Returns:
        dict with extracted key terms
    """
    client = anthropic.Anthropic(api_key=api_key or os.getenv("ANTHROPIC_API_KEY"))
    
    pdf_data = encode_pdf(pdf_path)
    
    prompt = """Extract key terms from this contract and output in JSON format:

```json
{
  "contract_type": "type of contract",
  "parties": {
    "party_a": {
      "name": "Party A name",
      "role": "Party A role (e.g., employer, seller)",
      "address": "address if available"
    },
    "party_b": {
      "name": "Party B name", 
      "role": "Party B role (e.g., employee, buyer)",
      "address": "address if available"
    }
  },
  "dates": {
    "effective_date": "effective date",
    "end_date": "termination date",
    "signing_date": "signing date"
  },
  "financial_terms": {
    "total_value": "total contract value",
    "payment_terms": "payment terms",
    "currency": "currency"
  },
  "key_clauses": [
    {
      "name": "clause name",
      "section": "section number",
      "summary": "clause summary",
      "original_text": "excerpt from original text"
    }
  ],
  "obligations": {
    "party_a": ["Party A obligation 1", "Party A obligation 2"],
    "party_b": ["Party B obligation 1", "Party B obligation 2"]
  },
  "termination": {
    "conditions": ["termination conditions"],
    "notice_period": "notice period",
    "consequences": "termination consequences"
  },
  "confidentiality": {
    "scope": "confidentiality scope",
    "duration": "confidentiality duration",
    "exceptions": ["exceptions"]
  },
  "dispute_resolution": {
    "method": "resolution method (litigation/arbitration)",
    "jurisdiction": "jurisdiction",
    "governing_law": "governing law"
  }
}
```

Requirements:
1. Use null for missing information
2. Use ISO date format (YYYY-MM-DD)
3. Include currency symbols with amounts
4. Return only JSON, no additional explanation

IMPORTANT: Extract values in the same language as the contract document.
"""
    
    response = client.messages.create(
        model=model,
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_data
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }
        ]
    )
    
    response_text = response.content[0].text
    
    # Try to parse JSON from response
    try:
        # Extract JSON from markdown code block if present
        json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', response_text)
        if json_match:
            json_str = json_match.group(1)
        else:
            json_str = response_text
        
        key_terms = json.loads(json_str)
    except json.JSONDecodeError:
        key_terms = {"raw_response": response_text, "parse_error": True}
    
    return {
        "status": "success",
        "key_terms": key_terms,
        "usage": {
            "input_tokens": response.usage.input_tokens,
            "output_tokens": response.usage.output_tokens
        }
    }


def generate_report(
    analysis: str,
    output_path: str,
    format: Literal["markdown", "docx"] = "markdown",
    title: str = "Contract Analysis Report"
) -> dict:
    """
    Generate a formatted report from analysis results.
    
    Args:
        analysis: The analysis text (markdown format)
        output_path: Path to save the report
        format: Output format ("markdown" or "docx")
        title: Report title
        
    Returns:
        dict with status and output path
    """
    output_path = Path(output_path)
    
    if format == "markdown":
        # Add header and metadata
        report_content = f"""# {title}

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Tool:** Contract Review Skill v0.1.0

---

{analysis}

---

*This report was generated by Contract Review Skill using Claude AI.*
*https://github.com/lijie420461340/contract-review-skill*
"""
        output_path.write_text(report_content, encoding="utf-8")
        return {"status": "success", "output_path": str(output_path), "format": "markdown"}
    
    elif format == "docx":
        if not DOCX_AVAILABLE:
            return {
                "status": "error",
                "message": "python-docx not installed. Run: pip install python-docx"
            }
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
        meta.add_run("Tool: Contract Review Skill v0.1.0").italic = True
        
        doc.add_paragraph()  # Spacer
        
        # Parse markdown and convert to DOCX
        lines = analysis.split('\n')
        current_list = None
        
        for line in lines:
            line = line.rstrip()
            
            # Headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # Numbered lists
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            
            # Bold text handling
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.strip('*')).bold = True
            
            # Regular paragraph
            elif line.strip():
                # Handle inline formatting
                p = doc.add_paragraph()
                
                # Simple bold/italic parsing
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        p.add_run(part[1:-1]).italic = True
                    else:
                        p.add_run(part)
            
            # Empty line = paragraph break
            elif not line.strip():
                pass  # Skip empty lines (already handled by paragraph structure)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run("This report was generated by Contract Review Skill using Claude AI.").italic = True
        footer.add_run("\nhttps://github.com/lijie420461340/contract-review-skill").italic = True
        
        doc.save(str(output_path))
        return {"status": "success", "output_path": str(output_path), "format": "docx"}
    
    else:
        return {"status": "error", "message": f"Unknown format: {format}"}
